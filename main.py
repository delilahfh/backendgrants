from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import json

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

class User(BaseModel):
    username: str
    password: str
    role: str  # "grantee" or "agency"

@app.post("/register")
def register(user: User):
    with open("users.json", "r+") as f:
        users = json.load(f)
        if user.username in users:
            raise HTTPException(status_code=400, detail="Username exists")
        users[user.username] = {"password": user.password, "role": user.role}
        f.seek(0)
        json.dump(users, f)
    return {"message": "User registered"}

@app.post("/login")
def login(user: User):
    with open("users.json") as f:
        users = json.load(f)
    if user.username not in users or users[user.username]["password"] != user.password:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    return {"role": users[user.username]["role"]}

from fastapi.responses import StreamingResponse
import io

@app.post("/narrative-proposal")
def narrative_proposal(data: dict):
    doc = Document()
    doc.add_heading("Narrative Proposal", 0)

    for key, value in data.items():
        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                             headers={"Content-Disposition": "attachment; filename=narrative_proposal.docx"})

from fastapi.responses import StreamingResponse
import io
from docx import Document

def generate_docx_from_data(data: dict, title: str) -> StreamingResponse:
    doc = Document()
    doc.add_heading(title, 0)
    for key, value in data.items():
        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={title.replace(' ', '_').lower()}.docx"}
    )

@app.post("/financial-proposal")
def financial_proposal(data: dict):
    return generate_docx_from_data(data, "Financial Proposal")

@app.post("/financial-amendment")
def financial_amendment(data: dict):
    return generate_docx_from_data(data, "Financial Amendment")

@app.post("/financial-report")
def financial_report(data: dict):
    return generate_docx_from_data(data, "Financial Report")

@app.post("/supporting-documents")
def supporting_documents(data: dict):
    return generate_docx_from_data(data, "Supporting Documents")

@app.post("/procurement-generator")
def procurement_generator(data: dict):
    return generate_docx_from_data(data, "Procurement Generator")

@app.post("/request-for-quotation")
def request_for_quotation(data: dict):
    return generate_docx_from_data(data, "Request For Quotation")

@app.post("/evaluation-report")
def evaluation_report(data: dict):
    return generate_docx_from_data(data, "Evaluation Report")

@app.post("/evaluation-annex")
def evaluation_annex(data: dict):
    return generate_docx_from_data(data, "Evaluation Annex")

@app.post("/procurement-contract")
def procurement_contract(data: dict):
    return generate_docx_from_data(data, "Procurement Contract")

@app.post("/delivery-certificate")
def delivery_certificate(data: dict):
    return generate_docx_from_data(data, "Delivery Certificate")

@app.post("/certificate-of-receipt")
def certificate_of_receipt(data: dict):
    return generate_docx_from_data(data, "Certificate Of Receipt")

@app.post("/fund-request")
def fund_request(data: dict):
    return generate_docx_from_data(data, "Fund Request")

@app.post("/grant-contract")
def grant_contract(data: dict):
    return generate_docx_from_data(data, "Grant Contract")