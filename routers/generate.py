from fastapi import APIRouter
from fastapi.responses import FileResponse
from pydantic import BaseModel
from datetime import datetime
import os
from docx import Document
from openpyxl import Workbook

router = APIRouter()
GEN_DIR = "generated"
os.makedirs(GEN_DIR, exist_ok=True)

class ProposalData(BaseModel):
    participant_code: str
    org_name: str
    business_name: str
    contact_info: str
    title: str
    description: str
    intervention: str
    impact: str
    funds_requested: float

@router.post("/narrative_proposal")
def generate_narrative(data: ProposalData):
    doc = Document()
    doc.add_heading("Narrative Proposal", 0)
    doc.add_paragraph(f"Participant Code: {data.participant_code}")
    doc.add_paragraph(f"Organization Name: {data.org_name}")
    doc.add_paragraph(f"Business Name: {data.business_name}")
    doc.add_paragraph(f"Contact Info: {data.contact_info}")
    doc.add_paragraph(f"Title: {data.title}")
    doc.add_paragraph(f"Description: {data.description}")
    doc.add_paragraph(f"Intervention: {data.intervention}")
    doc.add_paragraph(f"Expected Impact: {data.impact}")
    doc.add_paragraph(f"Requested Funds (USD): {data.funds_requested}")

    filename = f"narrative_{data.participant_code}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(GEN_DIR, filename)
    doc.save(filepath)
    return FileResponse(filepath, filename=filename)

class FinancialLine(BaseModel):
    category: str  # activity or supporting
    budget_line: str
    item: str
    quantity: int
    unit_cost: float
    provider: str

class FinancialData(BaseModel):
    participant_code: str
    org_name: str
    lines: list[FinancialLine]

@router.post("/financial_proposal")
def generate_financial(data: FinancialData):
    wb = Workbook()
    ws = wb.active
    ws.title = "Financial Proposal"
    ws.append(["Budget Line", "Item/Service", "Quantity", "Unit Cost", "Total Cost", "Provider"])

    total = 0
    for line in data.lines:
        cost = line.quantity * line.unit_cost
        total += cost
        ws.append([line.budget_line, line.item, line.quantity, line.unit_cost, cost, line.provider])

    ws.append([])
    ws.append(["", "", "", "Total:", total])

    filename = f"financial_{data.participant_code}_{datetime.now().strftime('%Y%m%d%H%M%S')}.xlsx"
    filepath = os.path.join(GEN_DIR, filename)
    wb.save(filepath)
    return FileResponse(filepath, filename=filename)